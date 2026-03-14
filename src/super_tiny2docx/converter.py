import io
import re

from bs4 import BeautifulSoup, Comment, NavigableString
from docx import Document as Docx
from docx.document import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.table import _Cell

from src.super_tiny2docx.doc_styles import ComputedStyle


class SuperTiny2Docx:
    """Convert to docx format with proper style inheritance."""
    style_manager = ComputedStyle

    def __init__(self, html_content: str):
        """
        Конструктор класса SuperTiny2Docx.
        :param html_content: Содержимое HTML.
        """
        self.html_content = html_content
        self.doc = None
        self.soup = None

    def convert(self, **kwargs) -> io.BytesIO:
        """Конвертирует HTML в DOCX"""

        self.doc = Docx()
        self._set_default_styles()
        self.soup = BeautifulSoup(self.html_content, "html.parser")

        main_content = self.soup.find("body")
        if not main_content:
            main_content = self.soup

        # Обрабатываем контент рекурсивно с вычислением стилей
        self._process_element(main_content, self.doc)
        # Сохраняем результат
        res_file = io.BytesIO()
        self.doc.save(res_file)
        res_file.seek(0)

        return res_file

    def _clear_document(self):
        """Очищает документ от содержимого, оставляя стили"""
        # Удаляем все параграфы из основного тела
        for paragraph in self.doc.paragraphs:
            p_element = paragraph._element
            p_element.getparent().remove(p_element)

        # Удаляем все таблицы из основного тела
        for table in self.doc.tables:
            t_element = table._element
            t_element.getparent().remove(t_element)

    def _set_default_styles(self):
        """Устанавливает стили по умолчанию для документа"""
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(14)

        # Устанавливаем отступ после параграфа по умолчанию = 0
        style.paragraph_format.space_after = Pt(0)

    def _process_paragraph(self, element, parent_docx_element, computed_style):
        """Обрабатывает параграф"""
        # Создаем новый параграф
        if isinstance(parent_docx_element, (_Cell, Document)):
            paragraph = parent_docx_element.add_paragraph()
        else:
            paragraph = parent_docx_element

        self._apply_paragraph_styles(paragraph, computed_style)

        # Обрабатываем дочерние элементы
        self._process_children(element, paragraph, computed_style)

    def _process_inline_container(self, element, parent_docx_element, computed_style):
        """Обрабатывает inline-контейнер (span, strong и т.д.) с сохранением пробелов"""

        # Проверяем, есть ли у элемента дети
        if not list(element.children):
            # Если нет детей, это может быть сам текст
            text = element.string
            if element.string and element.string.strip():
                self._process_text_with_context(
                    element.string, parent_docx_element, computed_style
                )
            return

        # Обрабатываем детей с сохранением контекста
        for child in element.children:
            if isinstance(child, NavigableString):
                # Это текст, обрабатываем его с текущими стилями
                text = str(child)
                if text:  # Сохраняем даже пустые строки? Нет, только если есть текст
                    # Но нам нужно сохранить пробелы, поэтому не делаем strip()
                    self._process_text_with_context(
                        text, parent_docx_element, computed_style
                    )
            else:
                # Это элемент, обрабатываем рекурсивно
                self._process_element(child, parent_docx_element, computed_style)

    def _process_text_with_context(self, text, parent_docx_element, computed_style):
        """Обрабатывает текст с учетом контекста и сохранением пробелов"""
        if not text:
            return

        # Находим или создаем параграф для текста
        if isinstance(parent_docx_element, _Cell):
            # Для ячеек таблицы создаем новый параграф, если нужно
            if not parent_docx_element.paragraphs:
                paragraph = parent_docx_element.add_paragraph()
            else:
                paragraph = parent_docx_element.paragraphs[-1]
        elif isinstance(parent_docx_element, Document):
            # Для корневого элемента
            paragraph = parent_docx_element.add_paragraph()
        else:
            # Для других случаев (уже параграф)
            paragraph = parent_docx_element

        # Применяем стили к параграфу (только если это новый параграф)
        if not paragraph.runs:
            self._apply_paragraph_styles(paragraph, computed_style)

        # Добавляем текст с сохранением пробелов и применяем стили
        run = paragraph.add_run(text)
        self._apply_run_styles(run, computed_style)

    def _process_line_break(self, parent_docx_element, computed_style):
        """Обрабатывает перенос строки"""
        if isinstance(parent_docx_element, _Cell):
            # Для ячеек таблицы добавляем новый параграф
            parent_docx_element.add_paragraph()
        elif isinstance(parent_docx_element, Document):
            parent_docx_element.add_paragraph()

    def _process_table(self, element, parent_docx_element, computed_style):
        """Обрабатывает таблицу"""
        # Ищем tbody
        tbody = element.find("tbody", recursive=False)
        if not tbody:
            # Если нет tbody, ищем строки непосредственно в таблице
            rows = element.find_all("tr", recursive=False)
            if not rows:
                return
            # Создаем фиктивный tbody для единообразия
            tbody = element

        rows = tbody.find_all("tr", recursive=False)
        if not rows:
            return

        # Определяем количество столбцов по первой строке
        first_row_cells = rows[0].find_all(["td", "th"], recursive=False)
        cols_count = len(first_row_cells)

        # Создаем таблицу в docx
        if isinstance(parent_docx_element, (_Cell, Document)):
            docx_table = parent_docx_element.add_table(rows=len(rows), cols=cols_count)
        else:
            docx_table = self.doc.add_table(rows=len(rows), cols=cols_count)

        # Применяем стили таблицы
        self._apply_table_styles(docx_table, computed_style)

        # Обрабатываем строки с передачей индексов через kwargs
        for row_index, row in enumerate(rows):
            # Передаем row_index в kwargs для обработки строки
            self._process_element(
                row,
                docx_table,
                computed_style,
                row_index=row_index,  # передаем индекс строки
            )

    def _process_row(
        self, element, parent_docx_element, computed_style, row_index=None
    ):
        """
        Обрабатывает строку таблицы с известным индексом
        :param element: HTML элемент строки (tr)
        :param parent_docx_element: DOCX таблица
        :param computed_style: вычисленные стили для строки
        :param row_index: индекс строки в таблице
        """
        # Находим все ячейки в строке
        cells = element.find_all(["td", "th"], recursive=False)

        # Обрабатываем каждую ячейку с передачей индексов
        for col_index, cell in enumerate(cells):
            # Передаем индексы строки и столбца в kwargs
            self._process_element(
                cell,
                parent_docx_element,
                computed_style,
                row_index=row_index,
                col_index=col_index,
            )

    def _process_cell(
        self,
        element,
        parent_docx_element,
        computed_style,
        row_index=None,
        col_index=None,
    ):
        """
        Обрабатывает ячейку таблицы с известными индексами
        :param element: HTML элемент ячейки (td или th)
        :param parent_docx_element: DOCX таблица
        :param computed_style: вычисленные стили для ячейки
        :param row_index: индекс строки
        :param col_index: индекс столбца
        """
        # if row_index is None or col_index is None:
        #     logger.warning(f"Cell indices not provided for {element.name}, skipping")
        #     return

        # Проверяем, что индексы в пределах таблицы
        if row_index >= len(parent_docx_element.rows) or col_index >= len(
            parent_docx_element.rows[row_index].cells
        ):
            raise IndexError(
                f"Cell indices out of range: row={row_index}, col={col_index}"
            )

        # Получаем ячейку в docx таблице
        docx_cell = parent_docx_element.cell(row_index, col_index)

        if list(element.children):
            # Очищаем ячейку от параграфов по умолчанию
            for paragraph in docx_cell.paragraphs:
                p_element = paragraph._element
                p_element.getparent().remove(p_element)

        # Применяем стили ячейки
        self._apply_cell_styles(docx_cell, computed_style)

        # Обрабатываем содержимое ячейки (без передачи индексов дальше)
        self._process_children(element, docx_cell, computed_style)

    def _process_list(self, element, parent_docx_element, computed_style):
        """Обрабатывает список"""
        items = element.find_all("li", recursive=False)

        for i, item in enumerate(items):
            # Создаем параграф для элемента списка
            if isinstance(parent_docx_element, (_Cell, Document)):
                paragraph = parent_docx_element.add_paragraph()
            else:
                paragraph = parent_docx_element

            if element.name == "ul":
                p_style = "List Bullet"
            else:
                p_style = "List Number"
            paragraph.style = p_style
            # run = paragraph.add_run(element.text.strip())
            # Добавляем маркер или номер
            # if element.name == "ol":
            #     prefix = f"{i + 1}. "
            # else:
            #     prefix = "• "

            # Обрабатываем содержимое элемента списка
            self._process_children(item, paragraph, computed_style)

    def _process_children(self, element, parent_docx_element, computed_style):
        """Обрабатывает дочерние элементы"""
        for child in element.children:
            self._process_element(child, parent_docx_element, computed_style)

    def _apply_paragraph_styles(self, paragraph, computed_style):
        """Применяет стили к параграфу"""
        # Устанавливаем отступ после по умолчанию = 0
        paragraph.paragraph_format.space_after = Pt(0)

        # Выравнивание
        paragraph.alignment = computed_style.get_text_align()

        # Отступы (margin)
        margin_top = computed_style.get("margin-top")
        if margin_top:
            num, unit = computed_style.get_numeric_value("margin-top", default_value=0)
            if unit == "pt":
                paragraph.paragraph_format.space_before = Pt(num)
            elif unit == "px":
                paragraph.paragraph_format.space_before = Pt(num * 0.75)

        margin_bottom = computed_style.get("margin-bottom")
        if margin_bottom:
            num, unit = computed_style.get_numeric_value(
                "margin-bottom", default_value=0
            )
            if unit == "pt":
                paragraph.paragraph_format.space_after = Pt(num)
            elif unit == "px":
                paragraph.paragraph_format.space_after = Pt(num * 0.75)

        margin_left = computed_style.get("margin-left")
        if margin_left:
            num, unit = computed_style.get_numeric_value("margin-left", default_value=0)
            if unit == "pt":
                paragraph.paragraph_format.left_indent = Pt(num)
            elif unit == "px":
                paragraph.paragraph_format.left_indent = Pt(num * 0.75)

        margin_right = computed_style.get("margin-right")
        if margin_right:
            num, unit = computed_style.get_numeric_value(
                "margin-right", default_value=0
            )
            if unit == "pt":
                paragraph.paragraph_format.right_indent = Pt(num)
            elif unit == "px":
                paragraph.paragraph_format.right_indent = Pt(num * 0.75)

        # Отступ первой строки
        text_indent = computed_style.get("text-indent")
        if text_indent:
            num, unit = computed_style.get_numeric_value("text-indent", default_value=0)
            if unit == "pt":
                paragraph.paragraph_format.first_line_indent = Pt(num)
            elif unit == "cm":
                paragraph.paragraph_format.first_line_indent = Cm(num)
            elif unit == "px":
                paragraph.paragraph_format.first_line_indent = Pt(num * 0.75)

    def _apply_run_styles(self, run, computed_style, parent_computed_style=None):
        """Применяет стили к run (тексту) с учетом наследования"""
        # Шрифт
        font_family = computed_style.get_font_family()
        run.font.name = (
            font_family.split(",")[0].strip().strip("'\"")
            if font_family
            else "Times New Roman"
        )

        # Размер шрифта - передаем родительский стиль для корректного вычисления процентов и em
        run.font.size = computed_style.get_font_size(parent_computed_style)

        # Начертание
        run.bold = computed_style.is_bold()
        run.italic = computed_style.is_italic()
        run.underline = computed_style.is_underlined()

        # Цвет текста
        color = computed_style.get_color()
        if color:
            try:
                run.font.color.rgb = self._parse_color(color)
            except:
                pass  # Игнорируем ошибки парсинга цвета

    def _apply_table_styles(self, table, computed_style):
        """Применяет стили к таблице"""
        # Границы
        if computed_style.get("border"):
            table.style = "Table Grid"

            # # Ширина таблицы
            # width = computed_style.get('width')
            # if width:
            #     if '%' in width:
            #         # Процентная ширина
            #         num, _ = computed_style.get_numeric_value('width',
            #                                                   default_value=100)
            #         self._set_table_width(table, width_type='pct',
            #                               value=int(num * 50))  # 100% = 5000
            #     else:
            #         # Абсолютная ширина
            #         num, unit = computed_style.get_numeric_value('width',
            #                                                      default_value=100)
            #         if unit == 'px':
            #             # Конвертируем px в twips (1px = 15 twips примерно)
            #             self._set_table_width(table, width_type='dxa',
            #                                   value=int(num * 15))
            #         elif unit == 'pt':
            #             self._set_table_width(table, width_type='dxa',
            #                                   value=int(
            #                                       num * 20))  # 1pt = 20 twips

        # Ширина таблицы
        self._set_table_width(table, width_type="pct", value=5000)
        table.autofit = False

    def _apply_cell_styles(self, cell, computed_style):
        """Применяет стили к ячейке"""
        # Вертикальное выравнивание
        valign = computed_style.get_vertical_align()
        self._set_cell_vertical_alignment(cell, valign)

        # Цвет фона
        bgcolor = computed_style.get_background_color()
        if bgcolor:
            self._set_cell_background_color(cell, bgcolor)

        # # Высота ячейки
        # height = computed_style.get('height')
        # if height:
        #     num, unit = computed_style.get_numeric_value('height', default_value=0)
        #     if unit == 'px':
        #         cell.height = Inches(num / 96)  # Приблизительно
        #     elif unit == 'pt':
        #         cell.height = Pt(num)

    def _set_table_width(self, table, width_type="pct", value=5000):
        """Устанавливает ширину таблицы"""
        tbl = table._tbl
        tblPr = tbl.tblPr

        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), str(value))
        tblW.set(qn("w:type"), width_type)

        # Удаляем существующий элемент ширины, если есть
        for elem in tblPr:
            if elem.tag == qn("w:tblW"):
                tblPr.remove(elem)

        tblPr.append(tblW)

    def _set_cell_vertical_alignment(self, cell, alignment):
        """Устанавливает вертикальное выравнивание для ячейки"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        vAlign = OxmlElement("w:vAlign")
        align_map = {"top": "top", "middle": "center", "bottom": "bottom"}
        vAlign.set(qn("w:val"), align_map.get(alignment, "top"))

        # Удаляем существующий элемент, если есть
        for elem in tcPr:
            if elem.tag == qn("w:vAlign"):
                tcPr.remove(elem)

        tcPr.append(vAlign)

    def _set_cell_background_color(self, cell, color):
        """Устанавливает цвет фона ячейки"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), self._parse_color(color))

        # Удаляем существующий элемент, если есть
        for elem in tcPr:
            if elem.tag == qn("w:shd"):
                tcPr.remove(elem)

        tcPr.append(shd)

    def _parse_color(self, color_string):
        """Парсит цвет в hex формат"""
        color_string = str(color_string).lower().strip()

        # Если это hex цвет
        if color_string.startswith("#"):
            return color_string[1:7]

        # Если это rgb
        rgb_match = re.search(r"rgb\((\d+),\s*(\d+),\s*(\d+)\)", color_string)
        if rgb_match:
            r, g, b = map(int, rgb_match.groups())
            return f"{r:02x}{g:02x}{b:02x}"

        # Словарь базовых цветов
        color_map = {
            "red": "FF0000",
            "green": "00FF00",
            "blue": "0000FF",
            "black": "000000",
            "white": "FFFFFF",
            "gray": "808080",
            "yellow": "FFFF00",
            "cyan": "00FFFF",
            "magenta": "FF00FF",
        }

        return color_map.get(color_string, "auto")

    def _process_element(
        self, element, parent_docx_element, parent_computed_style=None, **kwargs
    ):
        """
        Основной метод обработки элементов с вычислением стилей
        Теперь parent_computed_style содержит стили родителя с учетом всей цепочки наследования
        """
        element_name = element.name
        if isinstance(element, Comment):
            return
        if element.name is None:
            # Обрабатываем текст
            self._process_text(element, parent_docx_element, parent_computed_style)
            return

        # Вычисляем стили для текущего элемента с учетом родительских
        computed_style = self.style_manager(element, parent_computed_style)

        # Обрабатываем в зависимости от типа элемента
        if element.name == "table":
            self._process_table(element, parent_docx_element, computed_style)
        elif element.name == "tbody":
            # Для tbody просто обрабатываем детей, передавая стили от таблицы
            self._process_children(element, parent_docx_element, parent_computed_style)
        elif element.name == "tr":
            # Для строк таблицы получаем индекс из kwargs
            row_index = kwargs.get("row_index")
            self._process_row(element, parent_docx_element, computed_style, row_index)
        elif element.name in ["td", "th"]:
            # Для ячеек таблицы ожидаем, что индексы переданы через kwargs
            row_index = kwargs.get("row_index")
            col_index = kwargs.get("col_index")
            self._process_cell(
                element, parent_docx_element, computed_style, row_index, col_index
            )
        elif element.name in ["p", "div"]:
            self._process_paragraph(element, parent_docx_element, computed_style)
        elif element.name == "br":
            self._process_line_break(parent_docx_element, computed_style)
        elif element.name in ["ol", "ul"]:
            self._process_list(element, parent_docx_element, computed_style)
        elif element.name in ["span", "strong", "em", "b", "i", "u"]:
            self._process_inline_container(element, parent_docx_element, computed_style)
        else:
            # Для остальных элементов просто обрабатываем детей
            self._process_children(element, parent_docx_element, computed_style)

    def _process_text(self, text_node, parent_docx_element, computed_style):
        """Обрабатывает текстовый узел"""
        text = text_node.string
        if not text or not text.strip():
            return

        # Находим или создаем параграф для текста
        if isinstance(parent_docx_element, _Cell):
            # Для ячеек таблицы
            if not parent_docx_element.paragraphs:
                paragraph = parent_docx_element.add_paragraph()
            else:
                paragraph = parent_docx_element.paragraphs[-1]
        elif isinstance(parent_docx_element, Document):
            # Для корневого элемента
            paragraph = parent_docx_element.add_paragraph()
        else:
            # Для других случаев
            paragraph = parent_docx_element

        # Применяем стили к параграфу
        self._apply_paragraph_styles(paragraph, computed_style)

        # Добавляем текст и применяем стили текста
        run = paragraph.add_run(text.strip())
        self._apply_run_styles(
            run, computed_style, parent_computed_style=computed_style
        )


if __name__ == "__main__":
    html_content = """<p>Hello, world!</p>"""
    super_converter = SuperTiny2Docx(html_content)
    docx_file = super_converter.convert()

    with open("../../super_document.docx", "wb") as f:
        f.write(docx_file.read())
